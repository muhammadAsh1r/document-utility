# ================= Standard Library =================
import os
import re
import csv
import io
import unicodedata
from datetime import datetime

# ================= Document Libraries =================
import PyPDF2
from docx import Document


# ---------------- CONSTANTS ----------------
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".csv"}
HARMFUL_EXTENSIONS = {".exe", ".sh", ".bat", ".js", ".php"}


# ======================================================
#               FILENAME STANDARDIZATION
# ======================================================
def standardize_filename(filename: str) -> str:
    name, ext = os.path.splitext(filename)
    ext = ext.lower()

    name = re.sub(r"[^a-zA-Z0-9\s]", " ", name)
    name = re.sub(r"\s+", " ", name).strip()
    name = name.replace(" ", "_").upper()

    version_match = re.search(r"(\d+)$", name)
    version = version_match.group(1) if version_match else None

    if version:
        name = re.sub(r"_?\d+$", "", name)

    year = datetime.now().year
    parts = [name, str(year)]

    if version:
        parts.append(f"V{version}")

    return "_".join(parts) + ext


# ======================================================
#                  TEXT CLEANING
# ======================================================
def clean_text(text: str, remove_emojis=False) -> str:
    text = unicodedata.normalize("NFKC", text)

    replacements = {
        "“": '"', "”": '"',
        "‘": "'", "’": "'",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)

    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"[^\x00-\x7F]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()

    if remove_emojis:
        emoji_pattern = re.compile(
            "["
            "\U0001F600-\U0001F64F"
            "\U0001F300-\U0001F5FF"
            "\U0001F680-\U0001F6FF"
            "\U0001F700-\U0001F77F"
            "\U0001F900-\U0001F9FF"
            "\U0001FA00-\U0001FAFF"
            "]+",
            flags=re.UNICODE,
        )
        text = emoji_pattern.sub("", text)

    return text


# ======================================================
#                FILE VALIDATION
# ======================================================
def validate_file(uploaded_file):
    filename = uploaded_file.name.lower()
    ext = os.path.splitext(filename)[1]
    errors = []

    if ext in HARMFUL_EXTENSIONS or ext not in ALLOWED_EXTENSIONS:
        errors.append("Unsupported or harmful file extension")
        return errors

    if uploaded_file.size == 0:
        errors.append("File is empty")
        return errors

    try:
        if ext == ".pdf":
            reader = PyPDF2.PdfReader(uploaded_file)
            if not reader.pages:
                errors.append("PDF has no pages")

        elif ext == ".docx":
            doc = Document(uploaded_file)
            if not any(p.text.strip() for p in doc.paragraphs):
                errors.append("DOCX contains no readable text")

        elif ext == ".csv":
            decoded = uploaded_file.read().decode("utf-8", errors="ignore")
            rows = list(csv.reader(io.StringIO(decoded)))

            if not rows:
                errors.append("CSV file is empty")
                return errors

            header = rows[0]
            if not any(cell.strip() for cell in header):
                errors.append("CSV missing header")
                return errors

            expected_cols = len(header)
            for idx, row in enumerate(rows[1:], start=2):
                if not any(cell.strip() for cell in row):
                    continue
                if len(row) != expected_cols:
                    errors.append(
                        f"Unexpected number of columns at row {idx}: "
                        f"expected {expected_cols}, got {len(row)}"
                    )
                    break

    except Exception:
        errors.append("File is corrupted or unreadable")

    return errors


# ======================================================
#               METADATA EXTRACTION
# ======================================================
# ---------------- PDF ----------------
def extract_pdf_metadata(file):
    reader = PyPDF2.PdfReader(file)
    pages = reader.pages or []

    text_chunks = []
    for page in pages:
        page_text = page.extract_text()
        if page_text:
            text_chunks.append(page_text)

    text = "\n".join(text_chunks)
    meta = reader.metadata or {}

    title = meta.get("/Title")
    author = meta.get("/Author")

    # Fallbacks
    if not title:
        title = file.name

    return {
        "file_type": "pdf",
        "title": title,
        "author": author,
        "page_count": len(pages),
        "word_count": len(text.split()) if text else 0,
        "headings": extract_headings(text),
        "links": extract_links(text),
        "has_text": bool(text),   # 🔥 very useful for frontend
    }


# ---------------- DOCX ----------------
def extract_docx_metadata(file):
    doc = Document(file)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = " ".join(paragraphs)

    title = doc.core_properties.title or file.name
    author = doc.core_properties.author

    return {
        "file_type": "docx",
        "title": title,
        "author": author,
        "page_count": None,  # DOCX doesn't expose page count reliably
        "word_count": len(text.split()) if text else 0,
        "headings": [
            p.text for p in doc.paragraphs
            if p.style.name.startswith("Heading")
        ],
        "links": extract_links(text),
        "has_text": bool(text),
    }


# ---------------- Helpers ----------------
def extract_links(text: str):
    if not text:
        return []
    return re.findall(r"https?://\S+", text)


def extract_headings(text: str):
    if not text:
        return []
    return [
        line.strip()
        for line in text.split("\n")
        if line.isupper() and len(line.strip()) < 100
    ]
